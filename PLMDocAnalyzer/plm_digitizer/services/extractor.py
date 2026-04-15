"""
PLM Digitizer - Text Extraction Service
Handles PDF (searchable + scanned), DOCX, XLSX/XLS, Images, CSV/TXT
"""
import csv
import io
import logging
import os
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    file_path: str
    file_type: str
    raw_text: str
    extraction_method: str
    char_count: int
    success: bool
    error_message: Optional[str] = None


def extract_pdf(file_path: str) -> ExtractionResult:
    """Extract text from PDF — searchable first, fall back to OCR."""
    try:
        import pdfplumber

        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # Extract text
                page_text = page.extract_text() or ""
                text_parts.append(page_text)

                # Extract tables as markdown-style text
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        for row in table:
                            if row:
                                row_str = " | ".join(
                                    str(cell).strip() if cell else ""
                                    for cell in row
                                )
                                text_parts.append(row_str)

        full_text = "\n".join(text_parts).strip()

        # If very little text, try OCR
        if len(full_text) < 50:
            return extract_pdf_ocr(file_path)

        return ExtractionResult(
            file_path=file_path,
            file_type="PDF",
            raw_text=full_text,
            extraction_method="pdfplumber",
            char_count=len(full_text),
            success=True,
        )
    except Exception as e:
        # Try OCR fallback
        try:
            return extract_pdf_ocr(file_path)
        except Exception as e2:
            return ExtractionResult(
                file_path=file_path,
                file_type="PDF",
                raw_text="",
                extraction_method="failed",
                char_count=0,
                success=False,
                error_message=f"PDF extraction failed: {str(e2)}",
            )


def extract_pdf_ocr(file_path: str) -> ExtractionResult:
    """Extract text from scanned PDF using OCR."""
    try:
        from pdf2image import convert_from_path
        import pytesseract
        from PIL import Image, ImageFilter

        images = convert_from_path(file_path, dpi=200)
        text_parts = []

        for img in images:
            # Preprocess for better OCR
            gray = img.convert("L")
            # Apply threshold to binarize
            threshold = gray.point(lambda x: 0 if x < 128 else 255, "1")
            text = pytesseract.image_to_string(threshold, lang="eng")
            text_parts.append(text)

        full_text = "\n".join(text_parts).strip()
        return ExtractionResult(
            file_path=file_path,
            file_type="PDF",
            raw_text=full_text,
            extraction_method="ocr-pdf",
            char_count=len(full_text),
            success=True,
        )
    except Exception as e:
        raise RuntimeError(f"OCR failed: {str(e)}")


def extract_docx(file_path: str) -> ExtractionResult:
    """Extract text from DOCX."""
    try:
        from docx import Document

        doc = Document(file_path)
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_parts = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(row_parts))

        full_text = "\n".join(parts).strip()
        return ExtractionResult(
            file_path=file_path,
            file_type="DOCX",
            raw_text=full_text,
            extraction_method="python-docx",
            char_count=len(full_text),
            success=True,
        )
    except Exception as e:
        return ExtractionResult(
            file_path=file_path,
            file_type="DOCX",
            raw_text="",
            extraction_method="failed",
            char_count=0,
            success=False,
            error_message=str(e),
        )


def extract_excel(file_path: str, file_type: str = "XLSX") -> ExtractionResult:
    """Extract text from Excel files."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        parts = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            parts.append(f"[Sheet: {sheet_name}]")
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    val = cell.value
                    if val is not None:
                        row_values.append(str(val).strip())
                if any(row_values):
                    parts.append(" | ".join(row_values))

        wb.close()
        full_text = "\n".join(parts).strip()
        return ExtractionResult(
            file_path=file_path,
            file_type=file_type,
            raw_text=full_text,
            extraction_method="openpyxl",
            char_count=len(full_text),
            success=True,
        )
    except Exception as e:
        # Try XLS with xlrd fallback
        try:
            return extract_xls_fallback(file_path)
        except Exception:
            return ExtractionResult(
                file_path=file_path,
                file_type=file_type,
                raw_text="",
                extraction_method="failed",
                char_count=0,
                success=False,
                error_message=str(e),
            )


def extract_xls_fallback(file_path: str) -> ExtractionResult:
    """Fallback XLS extraction."""
    try:
        import xlrd
        wb = xlrd.open_workbook(file_path)
        parts = []
        for sheet_idx in range(wb.nsheets):
            sheet = wb.sheet_by_index(sheet_idx)
            parts.append(f"[Sheet: {sheet.name}]")
            for row_idx in range(sheet.nrows):
                row_values = [str(sheet.cell_value(row_idx, col)).strip()
                              for col in range(sheet.ncols)]
                if any(row_values):
                    parts.append(" | ".join(row_values))
        full_text = "\n".join(parts).strip()
        return ExtractionResult(
            file_path=file_path,
            file_type="XLS",
            raw_text=full_text,
            extraction_method="xlrd",
            char_count=len(full_text),
            success=True,
        )
    except Exception as e:
        raise RuntimeError(f"XLS extraction failed: {e}")


def extract_image(file_path: str, file_type: str = "PNG") -> ExtractionResult:
    """Extract text from image using OCR."""
    try:
        import pytesseract
        from PIL import Image, ImageFilter, ImageEnhance

        img = Image.open(file_path)
        # Convert to grayscale
        gray = img.convert("L")
        # Enhance contrast
        enhancer = ImageEnhance.Contrast(gray)
        enhanced = enhancer.enhance(2.0)
        # Binarize
        threshold = enhanced.point(lambda x: 0 if x < 140 else 255, "1")

        text = pytesseract.image_to_string(threshold, lang="eng")
        text = text.strip()

        return ExtractionResult(
            file_path=file_path,
            file_type=file_type,
            raw_text=text,
            extraction_method="pytesseract",
            char_count=len(text),
            success=True,
        )
    except Exception as e:
        return ExtractionResult(
            file_path=file_path,
            file_type=file_type,
            raw_text="",
            extraction_method="failed",
            char_count=0,
            success=False,
            error_message=str(e),
        )


def extract_csv_txt(file_path: str, file_type: str = "CSV") -> ExtractionResult:
    """Extract text from CSV or TXT file."""
    try:
        # Try to detect encoding
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
        raw = None
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc, errors="replace") as f:
                    raw = f.read()
                break
            except UnicodeDecodeError:
                continue

        if raw is None:
            with open(file_path, "rb") as f:
                raw = f.read().decode("latin-1", errors="replace")

        if file_type == "CSV":
            # Try to parse CSV and format as table
            try:
                dialect = csv.Sniffer().sniff(raw[:4096], delimiters=",;\t|")
                reader = csv.reader(io.StringIO(raw), dialect)
                rows = []
                for row in reader:
                    rows.append(" | ".join(str(c).strip() for c in row))
                    if len(rows) > 1000:  # Limit rows
                        break
                full_text = "\n".join(rows)
            except csv.Error:
                full_text = raw
        else:
            full_text = raw

        # Truncate very long texts
        if len(full_text) > 50000:
            full_text = full_text[:50000] + "\n[TRUNCATED]"

        return ExtractionResult(
            file_path=file_path,
            file_type=file_type,
            raw_text=full_text,
            extraction_method="csv-reader" if file_type == "CSV" else "text-read",
            char_count=len(full_text),
            success=True,
        )
    except Exception as e:
        return ExtractionResult(
            file_path=file_path,
            file_type=file_type,
            raw_text="",
            extraction_method="failed",
            char_count=0,
            success=False,
            error_message=str(e),
        )


def extract_file(file_path: str, file_type: str) -> ExtractionResult:
    """
    Main dispatch function — extracts text from any supported file.
    """
    start = time.time()

    file_type_upper = file_type.upper()

    try:
        if file_type_upper == "PDF":
            result = extract_pdf(file_path)
        elif file_type_upper in ("DOCX", "DOC"):
            result = extract_docx(file_path)
        elif file_type_upper in ("XLSX",):
            result = extract_excel(file_path, "XLSX")
        elif file_type_upper in ("XLS",):
            result = extract_excel(file_path, "XLS")
        elif file_type_upper in ("PNG", "JPG", "JPEG", "TIFF", "BMP"):
            result = extract_image(file_path, file_type_upper)
        elif file_type_upper in ("CSV",):
            result = extract_csv_txt(file_path, "CSV")
        elif file_type_upper in ("TXT",):
            result = extract_csv_txt(file_path, "TXT")
        else:
            result = ExtractionResult(
                file_path=file_path,
                file_type=file_type,
                raw_text="",
                extraction_method="unsupported",
                char_count=0,
                success=False,
                error_message=f"Unsupported file type: {file_type}",
            )
    except Exception as e:
        result = ExtractionResult(
            file_path=file_path,
            file_type=file_type,
            raw_text="",
            extraction_method="failed",
            char_count=0,
            success=False,
            error_message=str(e),
        )

    elapsed_ms = int((time.time() - start) * 1000)
    result.__dict__["processing_time_ms"] = elapsed_ms
    return result
